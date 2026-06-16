"""Parse a ``.docx`` brief into typed paragraphs (FR-01).

Defensive parsing: any structurally valid Office Open XML document is accepted;
a malformed or unreadable file fails loudly with a message naming it, and the
pipeline never partially processes (spec §2.1–2.2). The input document is opened
read-only here and is never modified — all downstream work happens on copies.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

import docx
from docx.document import Document as DocxDocument

_HEADING_RE = re.compile(r"^heading\s+(\d+)$", re.IGNORECASE)


class ParserError(Exception):
    """Raised when a ``.docx`` cannot be read or parsed."""


@dataclass(frozen=True)
class Paragraph:
    """One paragraph of the document body, with structural metadata."""

    index: int
    """Zero-based index of the paragraph in the document body."""

    text: str
    """The paragraph's plain text (runs concatenated, as python-docx reports)."""

    style_name: str
    """The paragraph's style name (e.g. ``Heading 1``, ``Body Text``)."""

    heading_level: int | None
    """The heading level (1, 2, …) for ``Heading N`` styles, else ``None``."""

    @property
    def is_heading(self) -> bool:
        """Whether this paragraph is a heading."""
        return self.heading_level is not None


@dataclass(frozen=True)
class ParsedDocument:
    """A parsed brief: the python-docx document plus typed paragraph metadata."""

    path: Path
    document: DocxDocument
    paragraphs: tuple[Paragraph, ...]


def _heading_level(style_name: str) -> int | None:
    """Return the heading level for a style name, or ``None`` if not a heading."""
    match = _HEADING_RE.match(style_name.strip())
    return int(match.group(1)) if match else None


def parse(path: Path) -> ParsedDocument:
    """Parse a ``.docx`` file into a :class:`ParsedDocument`.

    Args:
        path: Path to the input ``.docx`` brief.

    Returns:
        The parsed document with typed paragraph metadata.

    Raises:
        ParserError: If the file does not exist or is not a readable ``.docx``.
            The message names the file and the failure; no partial processing.
    """
    if not path.is_file():
        raise ParserError(f"input file '{path}' does not exist")
    try:
        document = docx.Document(str(path))
    except Exception as exc:  # python-docx raises a variety of errors
        raise ParserError(
            f"input file '{path}' is not a readable .docx document: {exc}"
        ) from exc

    paragraphs = tuple(
        Paragraph(
            index=index,
            text=para.text,
            style_name=para.style.name if para.style is not None else "",
            heading_level=_heading_level(para.style.name)
            if para.style is not None and para.style.name is not None
            else None,
        )
        for index, para in enumerate(document.paragraphs)
    )
    return ParsedDocument(path=path, document=document, paragraphs=paragraphs)
