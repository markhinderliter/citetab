"""Shared pytest fixtures and helpers for the toatool test suite.

This file is the single place that documents cross-cutting test concerns. In
this phase it provides:

- paths to the repo's versioned data (``rules``, ``profiles``, ``schemas``),
- loaded JSON Schemas for the registry and finding models,
- a builder for a minimal-but-complete valid :class:`CitationRegistry`,
- builders for the four derived fixtures (PRD §20), generated from the base
  briefs at test time so the derivation is inspectable and no opaque binary is
  committed,
- the fixed masked-field set used by the idempotency / baseline comparisons.
"""

from __future__ import annotations

import json
import shutil
from collections.abc import Callable
from pathlib import Path
from typing import Any

import docx
import pytest
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parent.parent
RULES_DIR = REPO_ROOT / "rules"
PROFILES_DIR = REPO_ROOT / "profiles"
SCHEMAS_DIR = REPO_ROOT / "schemas"
EXAMPLE_BRIEFS = REPO_ROOT / "examples" / "briefs"

#: Run-scoped fields masked before comparing serialized findings / reports to
#: baselines (PRD §20 "Masked non-deterministic fields"). Fixed and documented
#: here so every comparison uses the same set.
MASKED_FINDING_FIELDS = frozenset({"finding_id", "evaluated_at_utc", "engine_version"})
MASKED_RUN_METADATA_FIELDS = frozenset(
    {"input_sha256", "render_engine_version", "iteration_count"}
)


@pytest.fixture(scope="session")
def rules_dir() -> Path:
    """The repository ``rules`` directory."""
    return RULES_DIR


@pytest.fixture(scope="session")
def profiles_dir() -> Path:
    """The repository ``profiles`` directory."""
    return PROFILES_DIR


@pytest.fixture(scope="session")
def frap_profile_path() -> Path:
    """Path to the bundled FRAP court profile."""
    return PROFILES_DIR / "frap.yaml"


@pytest.fixture(scope="session")
def registry_schema() -> dict[str, Any]:
    """The parsed registry JSON Schema."""
    return json.loads(
        (SCHEMAS_DIR / "registry.schema.json").read_text(encoding="utf-8")
    )


@pytest.fixture(scope="session")
def finding_schema() -> dict[str, Any]:
    """The parsed finding JSON Schema."""
    return json.loads((SCHEMAS_DIR / "finding.schema.json").read_text(encoding="utf-8"))


_CARMODY_FULL = "Carmody v. Westfall Transit Auth., 512 F.3d 1042 (9th Cir. 2018)"


def _valid_registry_dict() -> dict[str, Any]:
    """Return a minimal-but-complete registry dict valid against the schema."""
    return {
        "authorities": [
            {
                "authority_id": "case::carmody-512-f3d-1042",
                "type": "case",
                "components": {
                    "reporter": "F.3d",
                    "volume": 512,
                    "first_page": 1042,
                    "court": "9th Cir.",
                    "year": 2018,
                },
                "display_full": _CARMODY_FULL,
                "display_short": "Carmody",
                "sort_key": "carmody v westfall transit auth",
                "group": "cases",
                "occurrences": [
                    {
                        "form": "full",
                        "raw_text": _CARMODY_FULL,
                        "paragraph_index": 12,
                        "char_span": [0, 61],
                        "confidence": "high",
                        "page": 2,
                    },
                    {
                        "form": "short",
                        "raw_text": "Carmody, 512 F.3d at 1045",
                        "paragraph_index": 20,
                        "char_span": [4, 29],
                        "confidence": "high",
                        "page": 3,
                        "pincite": "1045",
                    },
                ],
                "pages": [2, 3],
                "passim": False,
            }
        ],
        "run_metadata": {
            "engine_version": "0.1.0",
            "rule_pack_version": "1.0.0",
            "profile_id": "frap",
            "profile_version": "1.0.0",
            "render_engine": "LibreOffice",
            "render_engine_version": "24.2",
            "font_substitution_occurred": True,
            "iteration_count": 2,
            "converged": True,
            "input_sha256": "a" * 64,
            "font_substitutions": [
                {"original": "Times New Roman", "substitute": "Liberation Serif"}
            ],
        },
    }


@pytest.fixture()
def valid_registry_dict() -> Callable[[], dict[str, Any]]:
    """A factory returning a fresh valid registry dict each call.

    Returned as a factory (not the dict directly) so a test can mutate the
    result to construct invalid variants without affecting other tests.
    """
    return _valid_registry_dict


def _valid_finding_dict() -> dict[str, Any]:
    """Return a minimal-but-complete finding dict valid against the schema."""
    return {
        "finding_id": "f_0123456789ABCDEFGHJKMNPQRS",
        "rule_id": "TT-002",
        "rule_version": "1.0.0",
        "rule_pack_version": "1.0.0",
        "profile_id": "frap",
        "profile_version": "1.0.0",
        "severity": "info",
        "confidence": "high",
        "message": "Okafor was cited in the body but absent from the input TOA. Added.",
        "evidence": {
            "occurrence_references": [
                {
                    "paragraph_index": 45,
                    "char_span": [0, 47],
                    "page": 4,
                    "form": "full",
                    "excerpt": "United States v. Okafor, 891 F.3d 655 (9th Cir. 2019)",
                }
            ],
            "toa_entry_references": [],
            "computed_values": {"pages": [4, 6], "occurrence_count": 3},
        },
        "citations": [{"source": "FRAP", "section": "28(a)(2)"}],
        "subject": "United States v. Okafor, 891 F.3d 655 (9th Cir. 2019)",
        "evaluated_at_utc": "2026-06-09T12:00:00Z",
        "engine_version": "0.1.0",
    }


@pytest.fixture()
def valid_finding_dict() -> Callable[[], dict[str, Any]]:
    """A factory returning a fresh valid finding dict each call."""
    return _valid_finding_dict


@pytest.fixture()
def make_parsed() -> Callable[[list[tuple[str, str]]], Any]:
    """Factory building a synthetic ParsedDocument from (text, style) pairs.

    Heading level is derived from a ``Heading N`` style name. The ``document``
    field is left ``None`` (the pipeline stages under test read only paragraphs).
    """
    from toatool.pipeline.parser import Paragraph, ParsedDocument

    def _build(paras: list[tuple[str, str]]) -> Any:
        built = []
        for index, (text, style) in enumerate(paras):
            parts = style.split()
            level = (
                int(parts[1])
                if len(parts) == 2
                and parts[0].lower() == "heading"
                and parts[1].isdigit()
                else None
            )
            built.append(
                Paragraph(index=index, text=text, style_name=style, heading_level=level)
            )
        return ParsedDocument(
            path=Path("synthetic.docx"), document=None, paragraphs=tuple(built)
        )

    return _build


# --------------------------------------------------------------------------- #
# Derived fixtures (PRD §20) — built from the base briefs at test time.
# --------------------------------------------------------------------------- #

#: The phantom entry added to dirty for TT-004 (no matching body citation).
PHANTOM_ENTRY = "Nonexistent v. Authority, 999 F.9d 999   5"


def _style_object(document: Any, name: str) -> Any:
    """Return a style object used by some paragraph, by name, or ``None``.

    Reuses an existing paragraph's style object to sidestep python-docx's
    name-based lookup, which raises for built-in styles these fixtures reference
    but do not define under that exact name.
    """
    for para in document.paragraphs:
        if para.style is not None and para.style.name == name:
            return para.style
    return None


def _find(document: Any, text: str) -> Any:
    """Return the first paragraph whose stripped text equals ``text``."""
    return next(p for p in document.paragraphs if p.text.strip() == text)


def derive_dirty_plus_phantom(dest: Path) -> Path:
    """Copy the dirty brief and add a phantom (uncited) entry to its TOA → TT-004."""
    shutil.copy(EXAMPLE_BRIEFS / "dirty_motion_brief.docx", dest)
    document = docx.Document(str(dest))
    entry_style = _style_object(document, "Source Code")
    anchor = _find(document, "Statutes")
    para = anchor.insert_paragraph_before(PHANTOM_ENTRY)
    if entry_style is not None:
        para.style = entry_style
    document.save(str(dest))
    return dest


def derive_memo_no_marker(dest: Path) -> Path:
    """Copy the marker memo and delete its ``[[TOA]]`` marker → TT-005, suppressed."""
    shutil.copy(EXAMPLE_BRIEFS / "marker_trial_memo.docx", dest)
    document = docx.Document(str(dest))
    marker = _find(document, "[[TOA]]")
    marker._p.getparent().remove(marker._p)
    document.save(str(dest))
    return dest


def derive_brief_no_toa(dest: Path) -> Path:
    """Copy the clean brief and delete its whole TOA section → TT-005, suppressed.

    Removes the ``TABLE OF AUTHORITIES`` heading and every entry through the next
    ``Heading 1``, leaving no marker and no heading. Unlike the marker memo, this
    brief has a Carmody pinpoint that lands near a page boundary once the TOA is
    gone, so it regresses the no-placement measurement path (see QA Round 2).
    """
    shutil.copy(EXAMPLE_BRIEFS / "clean_appellate_brief.docx", dest)
    document = docx.Document(str(dest))
    paragraphs = document.paragraphs
    start = next(
        i
        for i, p in enumerate(paragraphs)
        if p.text.strip().upper() == "TABLE OF AUTHORITIES"
    )
    end = len(paragraphs)
    for j in range(start + 1, len(paragraphs)):
        if paragraphs[j].style.name == "Heading 1":
            end = j
            break
    for para in reversed(paragraphs[start:end]):
        para._p.getparent().remove(para._p)
    document.save(str(dest))
    return dest


def _split_with_page_break(paragraph: Any, marker: str) -> None:
    """Insert a hard page break immediately before ``marker`` in ``paragraph``.

    Rebuilds the paragraph's runs so a ``<w:br w:type="page"/>`` sits between the
    text before ``marker`` and ``marker`` onward. The break carries no text, so
    ``paragraph.text`` is unchanged and the citation still parses intact; only the
    *rendered* layout splits there.
    """
    text = paragraph.text
    index = text.find(marker)
    prefix, rest = text[:index], text[index:]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(prefix)
    break_run = paragraph.add_run()
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    break_run._element.append(page_break)
    paragraph.add_run(rest)


def derive_brief_unmeasured_occurrence(dest: Path) -> Path:
    """Copy the clean brief and split a pincite across a page boundary → TT-009.

    Injects a hard page break inside the Carmody pincite ``512 F.3d at 1047`` so
    its rendered text spans two physical pages. The locator joins pages with a
    newline, so that occurrence can no longer be matched and its page is
    unmeasured — exercising the emitting-path graceful-degradation path (QA Round
    3, C2-a). Deterministic and font/version-independent, unlike padding the body
    to shift pagination.
    """
    shutil.copy(EXAMPLE_BRIEFS / "clean_appellate_brief.docx", dest)
    document = docx.Document(str(dest))
    target = next(p for p in document.paragraphs if "512 F.3d at 1047" in p.text)
    _split_with_page_break(target, "1047")
    document.save(str(dest))
    return dest


def derive_dirty_plus_marker(dest: Path) -> Path:
    """Copy the dirty brief and add a ``[[TOA]]`` marker before the body → TT-006."""
    shutil.copy(EXAMPLE_BRIEFS / "dirty_motion_brief.docx", dest)
    document = docx.Document(str(dest))
    anchor = _find(document, "INTRODUCTION")
    anchor.insert_paragraph_before("[[TOA]]")
    document.save(str(dest))
    return dest


@pytest.fixture()
def dirty_plus_phantom(tmp_path: Path) -> Path:
    """Path to a freshly derived dirty-plus-phantom fixture."""
    return derive_dirty_plus_phantom(tmp_path / "dirty_plus_phantom.docx")


@pytest.fixture()
def memo_no_marker(tmp_path: Path) -> Path:
    """Path to a freshly derived memo-minus-marker fixture."""
    return derive_memo_no_marker(tmp_path / "memo_no_marker.docx")


@pytest.fixture()
def brief_no_toa(tmp_path: Path) -> Path:
    """Path to the clean brief with its entire TOA section removed (no marker)."""
    return derive_brief_no_toa(tmp_path / "brief_no_toa.docx")


@pytest.fixture()
def brief_unmeasured_occurrence(tmp_path: Path) -> Path:
    """Path to the clean brief with a pincite split across a page boundary."""
    return derive_brief_unmeasured_occurrence(
        tmp_path / "brief_unmeasured_occurrence.docx"
    )


@pytest.fixture()
def dirty_plus_marker(tmp_path: Path) -> Path:
    """Path to a freshly derived dirty-plus-marker fixture."""
    return derive_dirty_plus_marker(tmp_path / "dirty_plus_marker.docx")
