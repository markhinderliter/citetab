"""QA Round 3 — robustness & graceful-degradation harness.

Round 3 asks what a *user* gets on hostile or hard input. Unlike Round 2, it
drives the **real `citetab generate` CLI as a subprocess** — not the in-process
API — because the failure mode we care about (an uncaught exception) would crash
an in-process harness instead of being recorded. The bar (docs/QA_ROUND3.md):

- **reject** (Category 1, malformed/unreadable): exit 2, a one-line `error:`
  message to stderr, **nothing written**.
- **degrade** (Category 2, valid but hard; also valid non-briefs): no crash;
  the report is written (the `.docx` too, unless cleanly suppressed); any limit
  is disclosed via a finding.

**A Python traceback in stderr is always a FAIL**, in either mode.

Fixtures are *constructed* here (not committed binaries) so they stay
reproducible and inspectable. Some Category-2 cases are recorded as SKIP with a
rationale where faithful construction needs tooling python-docx lacks (Word
footnotes, section page-number restarts) or a real third-party document.

Usage::

    .venv-qa/bin/python scripts/qa_round3.py
"""

from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parent.parent
CLEAN = REPO_ROOT / "examples" / "briefs" / "clean_appellate_brief.docx"
RESULTS = REPO_ROOT / "docs" / "QA_ROUND3_RESULTS.md"

_TRACEBACK_MARKER = "Traceback (most recent call last)"
# OLE2 / CFB compound-file magic: the on-disk form of a legacy .doc AND of an
# encrypted OOXML package. python-docx hits the same "not a zip" path on both.
_OLE2_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")


# --------------------------------------------------------------------------- #
# Subprocess CLI runner — what the user actually sees.
# --------------------------------------------------------------------------- #
def _citetab() -> str:
    """Resolve the `citetab` console script next to the running interpreter."""
    candidate = Path(sys.executable).with_name("citetab")
    if not candidate.exists():
        raise FileNotFoundError(
            f"citetab console script not found at {candidate}; "
            "run the harness with the venv that has citetab installed"
        )
    return str(candidate)


@dataclass
class CliRun:
    """The observable result of one real `citetab generate` invocation."""

    exit_code: int
    stdout: str
    stderr: str
    files_written: list[str]

    @property
    def has_traceback(self) -> bool:
        """True if a Python traceback escaped to stderr (the cardinal sin)."""
        return _TRACEBACK_MARKER in self.stderr

    @property
    def report_written(self) -> bool:
        """True if a findings report was produced."""
        return any(name.endswith(".toa-report.md") for name in self.files_written)

    @property
    def docx_written(self) -> bool:
        """True if a regenerated .docx was produced."""
        return any(name.endswith(".toa.docx") for name in self.files_written)

    @property
    def stderr_tail(self) -> str:
        """The last non-empty stderr line (the user-facing error), if any."""
        lines = [ln for ln in self.stderr.strip().splitlines() if ln.strip()]
        return lines[-1] if lines else ""


def run_cli(brief: Path) -> CliRun:
    """Run `citetab generate <brief>` as a subprocess and capture what a user sees."""
    folder = brief.parent if brief.parent.exists() else REPO_ROOT
    stem = brief.stem
    before = {p.name for p in folder.iterdir()}
    proc = subprocess.run(  # noqa: S603 — fixed argv, local tool, QA harness
        [_citetab(), "generate", str(brief)],
        capture_output=True,
        text=True,
        cwd=str(folder),
    )
    after = {p.name for p in folder.iterdir()}
    written = sorted(
        name
        for name in (after - before)
        if name.startswith(stem) and (name.endswith(".docx") or name.endswith(".md"))
    )
    return CliRun(proc.returncode, proc.stdout, proc.stderr, written)


# --------------------------------------------------------------------------- #
# Construction helpers.
# --------------------------------------------------------------------------- #
def _clean_doc(dest: Path) -> docx.Document:  # type: ignore[name-defined]
    """Copy the clean fixture to ``dest`` and open it for editing."""
    shutil.copy(CLEAN, dest)
    return docx.Document(str(dest))


def _para_with(document: docx.Document, needle: str):  # type: ignore[name-defined]
    """Return the first paragraph containing ``needle``."""
    return next(p for p in document.paragraphs if needle in p.text)


def _insert_page_break_before_token(paragraph, token: str) -> None:
    """Rebuild ``paragraph`` with a hard page break inserted before ``token``."""
    text = paragraph.text
    index = text.find(token)
    prefix, rest = text[:index], text[index:]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(prefix)
    break_run = paragraph.add_run()
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    break_run._element.append(page_break)
    paragraph.add_run(rest)


# --- Category 1 (reject) builders ------------------------------------------ #
def build_missing(workdir: Path) -> Path:
    """A path that does not exist."""
    return workdir / "does_not_exist.docx"


def build_pdf_as_docx(workdir: Path) -> Path:
    """A PDF given a .docx name."""
    dest = workdir / "actually_a.docx"
    dest.write_bytes(b"%PDF-1.5\n%\xe2\xe3\xcf\xd3\n1 0 obj<<>>endobj\n")
    return dest


def build_truncated(workdir: Path) -> Path:
    """A real .docx truncated to its first half (broken zip)."""
    dest = workdir / "truncated.docx"
    data = CLEAN.read_bytes()
    dest.write_bytes(data[: len(data) // 2])
    return dest


def build_zero_byte(workdir: Path) -> Path:
    """A 0-byte file named .docx."""
    dest = workdir / "empty.docx"
    dest.write_bytes(b"")
    return dest


def build_legacy_doc(workdir: Path) -> Path:
    """A legacy .doc (OLE2 compound file) served as .docx."""
    dest = workdir / "legacy.docx"
    dest.write_bytes(_OLE2_MAGIC + b"\x00" * 512)
    return dest


def build_plain_text(workdir: Path) -> Path:
    """A plain-text file renamed .docx."""
    dest = workdir / "letter.docx"
    dest.write_text("Dear Court,\n\nThis is not a brief.\n", encoding="utf-8")
    return dest


def build_encrypted(workdir: Path) -> Path:
    """An encrypted OOXML package (OLE2/CFB on disk, like a real one)."""
    dest = workdir / "encrypted.docx"
    # Encrypted .docx files are OLE2 compound files containing an
    # "EncryptedPackage" stream; the OLE2 magic alone reproduces the same
    # not-a-zip rejection python-docx gives a real encrypted document.
    dest.write_bytes(_OLE2_MAGIC + b"\x00" * 1024)
    return dest


def build_directory(workdir: Path) -> Path:
    """A directory given a .docx name."""
    dest = workdir / "a_directory.docx"
    dest.mkdir()
    return dest


# --- Category 1 boundary + Category 2 (degrade) builders ------------------- #
def build_empty_valid(workdir: Path) -> Path:
    """A valid .docx with no text at all."""
    dest = workdir / "empty_valid.docx"
    docx.Document().save(str(dest))
    return dest


def build_non_brief(workdir: Path) -> Path:
    """A valid .docx of ordinary prose: no citations, no TOA, no marker."""
    dest = workdir / "non_brief.docx"
    document = docx.Document()
    document.add_paragraph("Memorandum")
    document.add_paragraph(
        "This document discusses the quarterly schedule and contains no legal "
        "citations and no table of authorities. It is not a brief."
    )
    document.save(str(dest))
    return dest


def build_c2a_boundary_split(workdir: Path) -> Path:
    """C2-a: a pincite split across a page boundary (the H1 fixture)."""
    dest = workdir / "c2a_boundary_split.docx"
    document = _clean_doc(dest)
    _insert_page_break_before_token(_para_with(document, "512 F.3d at 1047"), "1047")
    document.save(str(dest))
    return dest


def build_table_citation(workdir: Path) -> Path:
    """C2-c: a citation that lives only in a table cell (not a body paragraph)."""
    dest = workdir / "table_citation.docx"
    document = _clean_doc(dest)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Pemberton v. Quail, 701 F.3d 220 (9th Cir. 2014)"
    document.save(str(dest))
    return dest


def build_large_toa(workdir: Path) -> Path:
    """C2-d: a large brief whose regenerated TOA spans multiple pages."""
    dest = workdir / "large_toa.docx"
    document = _clean_doc(dest)
    anchor = _para_with(document, "The judgment should be reversed")
    for i in range(40):
        vol, page = 200 + i, 100 + i
        anchor.insert_paragraph_before(
            f"As the court explained in Marbrook{i:02d} v. Vance, {vol} F.3d "
            f"{page} (9th Cir. 20{i % 20:02d}), the standard is objective.",
            style=document.styles["Body Text"],
        )
    document.save(str(dest))
    return dest


def build_passim_heavy(workdir: Path) -> Path:
    """C2-e: one authority cited many times (passim territory)."""
    dest = workdir / "passim_heavy.docx"
    document = _clean_doc(dest)
    anchor = _para_with(document, "The judgment should be reversed")
    for n in range(1045, 1062):
        anchor.insert_paragraph_before(
            f"The court returned to the point. Carmody, 512 F.3d at {n}.",
            style=document.styles["Body Text"],
        )
    document.save(str(dest))
    return dest


def build_heading_variant(workdir: Path) -> Path:
    """C2-g: the TOA heading in a non-canonical casing the profile still accepts."""
    dest = workdir / "heading_variant.docx"
    document = _clean_doc(dest)
    heading = _para_with(document, "TABLE OF AUTHORITIES")
    for run in list(heading.runs):
        run._element.getparent().remove(run._element)
    heading.add_run("Table of Authorities")
    document.save(str(dest))
    return dest


def build_eyecite_hard(workdir: Path) -> Path:
    """C2-h: id./supra chains and a string citation."""
    dest = workdir / "eyecite_hard.docx"
    document = _clean_doc(dest)
    anchor = _para_with(document, "The judgment should be reversed")
    for line in (
        "The principle is settled. Carmody, supra, at 1051.",
        "Id. at 1052.",
        "See, e.g., Delgado, 487 U.S. at 220; Brunner, 344 P.3d at 905.",
    ):
        anchor.insert_paragraph_before(line, style=document.styles["Body Text"])
    document.save(str(dest))
    return dest


# --------------------------------------------------------------------------- #
# Cases.
# --------------------------------------------------------------------------- #
@dataclass
class Case:
    """One Round 3 case: how to build the input and how to judge the result."""

    slug: str
    title: str
    mode: str  # "reject" | "degrade"
    build: Callable[[Path], Path | None]
    skip_reason: str = ""


CASES: tuple[Case, ...] = (
    # Category 1 — reject (exit 2, nothing written).
    Case("C1-a", "Non-existent path", "reject", build_missing),
    Case("C1-b", "PDF renamed .docx", "reject", build_pdf_as_docx),
    Case("C1-c", "Truncated .docx (broken zip)", "reject", build_truncated),
    Case("C1-d", "0-byte .docx", "reject", build_zero_byte),
    Case("C1-e", "Legacy .doc (OLE2) as .docx", "reject", build_legacy_doc),
    Case("C1-f", "Plain text renamed .docx", "reject", build_plain_text),
    Case("C1-g", "Encrypted .docx (message quality)", "reject", build_encrypted),
    Case("C1-h", "Directory as input", "reject", build_directory),
    # Valid but text-empty: rejected with OCR guidance (exit 2), not TT-005.
    Case("C1-i", "Valid .docx, no readable text", "reject", build_empty_valid),
    # Valid prose without citations/TOA: degrades to TT-005 (exit 1), not exit 2.
    Case("C1-j", "Valid .docx, prose, no citations/TOA", "degrade", build_non_brief),
    # Category 2 — degrade (graceful on hard but valid input).
    Case(
        "C2-a",
        "Pincite split across a page boundary (H1)",
        "degrade",
        build_c2a_boundary_split,
    ),
    Case(
        "C2-b",
        "Footnote citations",
        "degrade",
        lambda _: None,
        "python-docx cannot author Word footnotes; needs raw-XML fixture or a "
        "real footnoted brief — deferred to a manual/real-document pass.",
    ),
    Case("C2-c", "Citation only in a table cell", "degrade", build_table_citation),
    Case(
        "C2-d",
        "Large brief, multi-page TOA (40+ authorities)",
        "degrade",
        build_large_toa,
    ),
    Case("C2-e", "Passim-heavy authority", "degrade", build_passim_heavy),
    Case(
        "C2-f",
        "Page-number restart (roman→arabic)",
        "degrade",
        lambda _: None,
        "section page-numbering restarts are low-fidelity via python-docx; "
        "deferred to a real-document pass.",
    ),
    Case(
        "C2-g",
        "TOA heading in a non-canonical casing",
        "degrade",
        build_heading_variant,
    ),
    Case(
        "C2-h", "eyecite-hard forms (id./supra/string)", "degrade", build_eyecite_hard
    ),
    Case(
        "C2-i",
        "Real public-domain brief",
        "degrade",
        lambda _: None,
        "out of automated scope; manual real-document check (do not commit "
        "third-party binaries).",
    ),
    Case(
        "C2-j",
        "True non-convergence (oscillation)",
        "degrade",
        lambda _: None,
        "impractical/non-deterministic to construct; loop logic is covered by "
        "the mocked-measurer unit test test_tt007_convergence.",
    ),
)


@dataclass
class CaseResult:
    """The verdict for one case."""

    case: Case
    status: str  # "PASS" | "FAIL" | "SKIP"
    detail: str
    message: str = ""
    files: list[str] = field(default_factory=list)


def evaluate(case: Case, workdir: Path) -> CaseResult:
    """Build the case input, run the real CLI, and judge it against the bar."""
    path = case.build(workdir)
    if path is None:
        return CaseResult(case, "SKIP", case.skip_reason)

    run = run_cli(path)
    if run.has_traceback:
        return CaseResult(
            case,
            "FAIL",
            f"traceback escaped to stderr (exit {run.exit_code})",
            run.stderr_tail,
            run.files_written,
        )

    if case.mode == "reject":
        ok = run.exit_code == 2 and not run.files_written
        detail = (
            f"exit={run.exit_code} (want 2); "
            f"files={run.files_written or 'none'} (want none)"
        )
        return CaseResult(
            case, "PASS" if ok else "FAIL", detail, run.stderr_tail, run.files_written
        )

    # degrade: no crash, a report was written, and the run did not error out
    # in an uncontrolled way (exit 0 = warnings/info; exit 1 = error finding or
    # a cleanly suppressed .docx).
    ok = run.report_written and run.exit_code in (0, 1)
    detail = (
        f"exit={run.exit_code}; report={'yes' if run.report_written else 'NO'}; "
        f"docx={'yes' if run.docx_written else 'suppressed/none'}; "
        f"files={run.files_written or 'none'}"
    )
    return CaseResult(
        case, "PASS" if ok else "FAIL", detail, run.stderr_tail, run.files_written
    )


# --------------------------------------------------------------------------- #
# Reporting.
# --------------------------------------------------------------------------- #
def write_results(results: list[CaseResult]) -> None:
    """Write the Round 3 results report."""
    passed = sum(1 for r in results if r.status == "PASS")
    failed = sum(1 for r in results if r.status == "FAIL")
    skipped = sum(1 for r in results if r.status == "SKIP")
    ran = passed + failed

    lines = [
        "# QA Round 3 — robustness results",
        "",
        "Generated by `scripts/qa_round3.py`. Method and the full case list: see "
        "[QA_ROUND3.md](QA_ROUND3.md). The harness drives the real `citetab "
        "generate` CLI as a subprocess; a Python traceback in stderr is a FAIL.",
        "",
        f"**{passed}/{ran} executed cases passed"
        f"{'' if not failed else f' — {failed} FAILED'}; {skipped} skipped "
        f"(construction deferred — see each).**",
        "",
        "| Case | Mode | Status | Detail |",
        "|------|------|--------|--------|",
    ]
    for r in results:
        detail = r.detail if r.status != "SKIP" else f"skipped — {r.detail}"
        lines.append(f"| {r.case.slug} | {r.case.mode} | {r.status} | {detail} |")

    rejects = [r for r in results if r.case.mode == "reject" and r.status != "SKIP"]
    lines += ["", "## Category 1 — message quality (reject cases)", ""]
    for r in rejects:
        msg = r.message or "(no stderr)"
        lines.append(f"- **{r.case.slug}** ({r.case.title}): `{msg}`")
    lines += [
        "",
        "Note: an encrypted `.docx` and a legacy `.doc` are both OLE2 compound "
        "files on disk, so the tool rejects them with the same generic "
        '"not a readable .docx" message — honest and human, but it cannot '
        "single out *encryption* as the cause without sniffing the OLE2 magic. "
        "A more specific hint is a possible (cosmetic) enhancement, not a "
        "robustness gap.",
        "",
        "## Skips (deferred construction)",
        "",
    ]
    for r in results:
        if r.status == "SKIP":
            lines.append(f"- **{r.case.slug}** ({r.case.title}): {r.detail}")
    lines.append("")
    RESULTS.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    """Run every Round 3 case and report. Non-zero exit iff a case FAILED."""
    results: list[CaseResult] = []
    with tempfile.TemporaryDirectory() as td:
        for case in CASES:
            case_dir = Path(td) / case.slug
            case_dir.mkdir()
            result = evaluate(case, case_dir)
            results.append(result)
            print(f"[{result.status:4}] {case.slug:5} {case.title}: {result.detail}")
    write_results(results)
    failed = [r for r in results if r.status == "FAIL"]
    passed = sum(1 for r in results if r.status == "PASS")
    ran = passed + len(failed)
    print(f"\n{passed}/{ran} passed, {len(failed)} failed — summary: {RESULTS}")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
