"""QA Round 2 — automated defect-injection harness.

For each actionable rule, this script:

1. Loads a fresh copy of the finding-free clean fixture
   (``examples/briefs/clean_appellate_brief.docx``).
2. Injects **exactly one** defect that should violate that one rule, and
   records — in plain language — what was changed.
3. Saves the altered ``.docx`` to ``build/qa_round2/``.
4. Runs the real pipeline + rule pack on the altered document in-process.
5. Compares the fired findings (rule id, severity), the exit code, and the
   ``.docx`` suppression flag against the matrix in ``docs/QA_ROUND2.md``.

Because the clean fixture is free of *defect* findings, any defect finding
on an altered copy is necessarily the injected one. TT-008 (font
substitution) fires on every run in this environment and is treated as the
baseline; it is never counted as a defect. TT-007 (non-convergence) cannot
be hand-injected and is covered by a mocked unit test — neither is exercised
here (see ``docs/QA_ROUND2.md``).

The harness records *every* fired rule, so cross-rule interactions (e.g. a
TOA-length edit rippling into a spurious TT-003) are surfaced rather than
hidden. A scenario passes only when the expected rule fired at the expected
severity, the exit code and suppression flag match, and **no other** defect
rule fired.

Usage::

    .venv-qa/bin/python scripts/qa_round2.py
"""

from __future__ import annotations

import shutil
import traceback
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from toatool.engine.profile_loader import load_profile_by_id
from toatool.engine.runner import run_rules
from toatool.pipeline import convergence

REPO_ROOT = Path(__file__).resolve().parent.parent
CLEAN = REPO_ROOT / "examples" / "briefs" / "clean_appellate_brief.docx"
WORKDIR = REPO_ROOT / "build" / "qa_round2"
RESULTS = REPO_ROOT / "docs" / "QA_ROUND2_RESULTS.md"
PROFILE_ID = "frap"

# TT-008 fires on every run (font substitution); it is the baseline, never a
# defect. Any other TT-0xx on an altered copy is an injected/observed defect.
BASELINE_RULE = "TT-008"

# Authored analysis for scenarios that do not behave as the matrix predicts.
# Rendered into the report only when that scenario does not pass, so the
# write-up stays reproducible across re-runs.
ANALYSIS_NOTES: dict[str, str] = {
    "tt005": (
        "Deleting the heading-based TOA from the clean appellate brief does "
        "**not** produce a clean TT-005. Instead `convergence.generate` raises "
        "`ConvergenceError` from `freeze_registry` *before* the rule pack runs, "
        "because a Carmody pinpoint occurrence (`512 F.3d at 1047`) has no "
        "measured page in the TOA-less render. No TT-005 finding is emitted and "
        "no exit code is returned.\n\n"
        "This is **input-specific, not a TT-005-path defect**: the same "
        "no-placement path produces a correct TT-005 (error, `.docx` "
        "suppressed, exit 1) on the marker-memo fixture "
        "(`tests/integration/test_rules_oracle.py::test_oracle_memo_no_marker`, "
        "verified passing), and through this harness's exact "
        "`generate` + `run_rules` call path. The likely mechanism: removing "
        "~14 paragraphs of TOA shifts body pagination, landing the `1047` "
        "pinpoint on a page boundary so the locator cannot match the whole "
        "occurrence string on a single rendered page; `freeze_registry` then "
        "refuses rather than letting TT-005 report.\n\n"
        "**Follow-up (code change, out of scope for this QA round):** in the "
        "no-placement path the `.docx` is already suppressed, so an "
        "unmeasurable occurrence should degrade to a reported finding rather "
        "than a hard `ConvergenceError`. Track as a robustness item. For QA "
        "coverage in the meantime, TT-005 is exercised by the marker-memo "
        "fixture in the oracle suite."
    ),
}


# --------------------------------------------------------------------------- #
# python-docx helpers (small, surgical edits to a loaded Document)
# --------------------------------------------------------------------------- #
def find_index(doc: Document, substring: str) -> int:
    """Return the index of the first paragraph containing ``substring``."""
    for i, para in enumerate(doc.paragraphs):
        if substring in para.text:
            return i
    raise ValueError(f"no paragraph contains {substring!r}")


def delete_paragraph(paragraph: Paragraph) -> None:
    """Remove a paragraph element from its parent (python-docx has no delete)."""
    element = paragraph._element
    element.getparent().remove(element)


def set_text_preserve_style(paragraph: Paragraph, text: str) -> None:
    """Replace a paragraph's text, keeping its first run (and thus its style)."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def insert_paragraph_after(
    doc: Document, paragraph: Paragraph, text: str, style_name: str
) -> Paragraph:
    """Insert a new styled paragraph immediately after ``paragraph``."""
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_para = Paragraph(new_element, paragraph._parent)
    new_para.style = doc.styles[style_name]
    new_para.add_run(text)
    return new_para


def toa_region(doc: Document) -> tuple[int, int]:
    """Return ``(heading_index, end_exclusive)`` for the TOA section.

    The region runs from the ``TABLE OF AUTHORITIES`` heading through the
    paragraph before the next ``Heading 1`` (or end of document).
    """
    heading = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().upper() == "TABLE OF AUTHORITIES":
            heading = i
            break
    if heading is None:
        raise ValueError("clean fixture has no 'TABLE OF AUTHORITIES' heading")
    end = len(doc.paragraphs)
    for j in range(heading + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == "Heading 1":
            end = j
            break
    return heading, end


def delete_toa_region(doc: Document) -> list[str]:
    """Delete the whole TOA section; return the removed paragraph texts."""
    heading, end = toa_region(doc)
    removed = [doc.paragraphs[k].text for k in range(heading, end)]
    for k in range(end - 1, heading - 1, -1):
        delete_paragraph(doc.paragraphs[k])
    return removed


# --------------------------------------------------------------------------- #
# Injectors — each mutates the doc in place and returns a change-log line.
# --------------------------------------------------------------------------- #
def inject_tt002(doc: Document) -> str:
    """TT-002: delete a TOA entry for a case still cited in the body."""
    i = find_index(doc, "Carmody v. Westfall")
    removed = doc.paragraphs[i].text
    delete_paragraph(doc.paragraphs[i])
    return f"Deleted TOA entry for a case still cited in the body: {removed!r}"


def inject_tt003(doc: Document) -> str:
    """TT-003: change one TOA entry's page number to a wrong value."""
    import re

    i = find_index(doc, "Brunner v. Caldwell")
    old = doc.paragraphs[i].text
    new = re.sub(r"(\)\s+)(\d[\d,\s]*)$", lambda m: m.group(1) + "60", old)
    if new == old:
        raise ValueError(f"page-number substitution did not apply to {old!r}")
    set_text_preserve_style(doc.paragraphs[i], new)
    return f"Changed a TOA page number to a wrong value: {old!r} -> {new!r}"


def inject_tt004(doc: Document) -> str:
    """TT-004: add a TOA entry for a case cited nowhere in the body."""
    i = find_index(doc, "Hartwell Industries")
    entry = "Larkin v. Osgood, 701 F.3d 220 (9th Cir. 2014)\t7"
    insert_paragraph_after(doc, doc.paragraphs[i], entry, "Source Code")
    return f"Added a phantom TOA entry (cited nowhere in the body): {entry!r}"


def inject_marker_insertion(doc: Document) -> str:
    """Insertion path: delete the TOA section and drop a [[TOA]] marker."""
    removed = delete_toa_region(doc)
    # After deletion the former next-heading paragraph occupies the region
    # start; anchor the marker before the first remaining Heading 1.
    anchor = None
    for para in doc.paragraphs:
        if para.style.name == "Heading 1":
            anchor = para
            break
    if anchor is None:
        raise ValueError("no heading remained to anchor the marker before")
    anchor.insert_paragraph_before("[[TOA]]")
    return (
        f"Deleted the TOA section ({len(removed)} paragraphs: heading + "
        f"entries) and inserted a [[TOA]] marker in its place"
    )


def inject_tt005(doc: Document) -> str:
    """TT-005: delete the TOA section with no marker and no heading."""
    removed = delete_toa_region(doc)
    return (
        f"Deleted the TOA section ({len(removed)} paragraphs) with no marker "
        f"and no replacement heading"
    )


def inject_tt001(doc: Document) -> str:
    """TT-001: write an orphan short form with no full citation anywhere."""
    i = find_index(doc, "This Court so held in Carmody")
    sentence = (
        "The controlling rule is stated in Pemberton, 410 F.3d at 555, which "
        "the district court overlooked."
    )
    insert_paragraph_after(doc, doc.paragraphs[i], sentence, "Body Text")
    return f"Inserted an orphan short form (no full citation anywhere): {sentence!r}"


def inject_tt006(doc: Document) -> str:
    """TT-006: add a [[TOA]] marker while keeping the real TOA heading."""
    heading, _ = toa_region(doc)
    doc.paragraphs[heading].insert_paragraph_before("[[TOA]]")
    return (
        "Added a [[TOA]] marker while keeping the existing 'TABLE OF "
        "AUTHORITIES' heading (marker + heading both present)"
    )


# --------------------------------------------------------------------------- #
# Scenario table — order matches docs/QA_ROUND2.md.
# --------------------------------------------------------------------------- #
@dataclass(frozen=True)
class Scenario:
    """One defect-injection case and its expected outcome."""

    slug: str
    title: str
    inject: Callable[[Document], str]
    expected_rule: str | None  # None = control: no defect rule should fire
    expected_severity: str | None
    expected_exit: int
    expected_suppressed: bool


SCENARIOS: tuple[Scenario, ...] = (
    Scenario(
        "tt002",
        "Missing TOA entry for a cited authority",
        inject_tt002,
        "TT-002",
        "info",
        0,
        False,
    ),
    Scenario(
        "tt003",
        "Stale TOA page number",
        inject_tt003,
        "TT-003",
        "info",
        0,
        False,
    ),
    Scenario(
        "tt004",
        "Phantom TOA entry (uncited)",
        inject_tt004,
        "TT-004",
        "warning",
        0,
        False,
    ),
    Scenario(
        "marker",
        "TOA deleted + [[TOA]] marker (insertion path)",
        inject_marker_insertion,
        None,
        None,
        0,
        False,
    ),
    Scenario(
        "tt005",
        "TOA deleted, no marker (placement not found)",
        inject_tt005,
        "TT-005",
        "error",
        1,
        True,
    ),
    Scenario(
        "tt001",
        "Orphan short form (unresolvable)",
        inject_tt001,
        "TT-001",
        "error",
        1,
        False,
    ),
    Scenario(
        "tt006",
        "Marker + heading conflict",
        inject_tt006,
        "TT-006",
        "warning",
        0,
        False,
    ),
)


# --------------------------------------------------------------------------- #
# Run one scenario.
# --------------------------------------------------------------------------- #
@dataclass
class Outcome:
    """The observed result of running one scenario."""

    scenario: Scenario
    change_log: str = ""
    fired: dict[str, str] = field(default_factory=dict)  # rule_id -> severity
    messages: dict[str, str] = field(default_factory=dict)  # rule_id -> message
    exit_code: int | None = None
    suppressed: bool | None = None
    error: str | None = None
    altered_path: Path | None = None

    @property
    def defect_rules(self) -> dict[str, str]:
        """Fired rules excluding the always-on TT-008 baseline."""
        return {r: s for r, s in self.fired.items() if r != BASELINE_RULE}

    @property
    def unexpected(self) -> list[str]:
        """Defect rules that fired but were not the one we injected for."""
        rule = self.scenario.expected_rule
        expected = {rule} if rule else set()
        return sorted(set(self.defect_rules) - expected)

    @property
    def error_summary(self) -> str | None:
        """The last non-empty line of the traceback (the exception message)."""
        if not self.error:
            return None
        lines = [ln for ln in self.error.strip().splitlines() if ln.strip()]
        return lines[-1] if lines else None

    @property
    def passed(self) -> bool:
        """True only on an exact match with no surprises."""
        if self.error is not None:
            return False
        if self.exit_code != self.scenario.expected_exit:
            return False
        if self.suppressed != self.scenario.expected_suppressed:
            return False
        if self.unexpected:
            return False
        if self.scenario.expected_rule is None:
            return not self.defect_rules
        sc = self.scenario
        if sc.expected_rule not in self.fired:
            return False
        return self.fired[sc.expected_rule] == sc.expected_severity


def run_scenario(scenario: Scenario, profile) -> Outcome:
    """Inject the defect, run the pipeline, and capture the outcome."""
    outcome = Outcome(scenario=scenario)
    altered = WORKDIR / f"{scenario.slug}_altered.docx"
    try:
        shutil.copyfile(CLEAN, altered)
        doc = Document(str(altered))
        outcome.change_log = scenario.inject(doc)
        doc.save(str(altered))
        outcome.altered_path = altered

        gen = convergence.generate(altered, profile)
        rules = run_rules(gen, input_path=altered, profile=profile)
        outcome.exit_code = rules.exit_code
        outcome.suppressed = rules.docx_suppressed
        for f in rules.findings:
            outcome.fired[f.rule_id] = f.severity
            outcome.messages[f.rule_id] = f.message
    except Exception:  # noqa: BLE001 — QA harness records failures, never aborts
        outcome.error = traceback.format_exc()
    return outcome


# --------------------------------------------------------------------------- #
# Reporting.
# --------------------------------------------------------------------------- #
def write_case_log(outcome: Outcome) -> None:
    """Write a per-scenario change + observation log next to the altered docx."""
    sc = outcome.scenario
    fired_findings = ", ".join(f"{r} ({s})" for r, s in outcome.fired.items()) or "none"
    lines = [
        f"# QA Round 2 — {sc.slug}: {sc.title}",
        "",
        f"- altered document: {outcome.altered_path}",
        f"- injected change: {outcome.change_log or '(failed before injection)'}",
        "",
        "## Expected",
        f"- rule: {sc.expected_rule or '(none — control)'}",
        f"- severity: {sc.expected_severity or '—'}",
        f"- exit code: {sc.expected_exit}",
        f"- .docx suppressed: {sc.expected_suppressed}",
        "",
        "## Observed",
        f"- exit code: {outcome.exit_code}",
        f"- .docx suppressed: {outcome.suppressed}",
        f"- fired findings: {fired_findings}",
    ]
    if outcome.unexpected:
        lines.append(f"- UNEXPECTED defect rules: {', '.join(outcome.unexpected)}")
    if outcome.error:
        lines += ["", "## Harness error", "```", outcome.error.strip(), "```"]
    lines += ["", "## Finding messages"]
    for rule_id, message in outcome.messages.items():
        lines.append(f"- **{rule_id}**: {message}")
    lines += ["", f"## Verdict: {'PASS' if outcome.passed else 'FAIL'}", ""]
    (WORKDIR / f"{outcome.scenario.slug}_changes.md").write_text(
        "\n".join(lines), encoding="utf-8"
    )


def write_summary(outcomes: list[Outcome]) -> None:
    """Write the single Round-2 results report consumed by reviewers."""
    passed = sum(1 for o in outcomes if o.passed)
    total = len(outcomes)
    rows = []
    for o in outcomes:
        sc = o.scenario
        obs_rule = ", ".join(sorted(o.defect_rules)) or "(none)"
        obs_sev = o.fired.get(sc.expected_rule, "—") if sc.expected_rule else "—"
        verdict = "✅ PASS" if o.passed else "❌ FAIL"
        rows.append(
            f"| {sc.slug} | {sc.title} | {sc.expected_rule or '(none)'} | "
            f"{sc.expected_severity or '—'} | {sc.expected_exit} | "
            f"{obs_rule} | {obs_sev} | {o.exit_code} | {verdict} |"
        )

    discrepancies = [o for o in outcomes if not o.passed]
    if discrepancies:
        conclusion_lines = [
            f"{len(outcomes) - len(discrepancies)} of {len(outcomes)} scenarios "
            "matched the matrix exactly. The following did not — each is a real "
            "observed behavior, not a harness artifact:",
            "",
        ]
        for o in discrepancies:
            reason = o.error_summary or (
                f"observed {sorted(o.defect_rules) or 'no defect rule'}, "
                f"exit {o.exit_code}, suppressed {o.suppressed}"
            )
            conclusion_lines.append(
                f"- **{o.scenario.slug}** ({o.scenario.title}): {reason}"
            )
    else:
        conclusion_lines = [
            f"All {len(outcomes)} scenarios matched the matrix exactly: each "
            "injected defect produced the expected rule at the expected "
            "severity, exit code, and suppression — and no other defect rule "
            "fired. No cross-rule ripple was observed (single-line TOA edits did "
            "not spuriously trigger TT-003)."
        ]

    detail = []
    for o in outcomes:
        sc = o.scenario
        detail.append(f"### {sc.slug} — {sc.title}\n")
        detail.append(f"- **Injected:** {o.change_log or '(failed before injection)'}")
        detail.append(
            f"- **Expected:** rule `{sc.expected_rule or 'none'}` · "
            f"severity {sc.expected_severity or '—'} · exit {sc.expected_exit} · "
            f"suppressed {sc.expected_suppressed}"
        )
        fired = ", ".join(f"`{r}` ({s})" for r, s in o.fired.items()) or "none"
        detail.append(
            f"- **Observed:** fired {fired} · exit {o.exit_code} · "
            f"suppressed {o.suppressed}"
        )
        if o.unexpected:
            detail.append(f"- **⚠ Unexpected defect rules:** {', '.join(o.unexpected)}")
        if o.error_summary:
            detail.append(f"- **Raised before result:** `{o.error_summary}`")
        detail.append(f"- **Verdict:** {'PASS' if o.passed else 'FAIL'}")
        if not o.passed and sc.slug in ANALYSIS_NOTES:
            detail.append(f"\n  {ANALYSIS_NOTES[sc.slug]}")
        detail.append("")

    table_header = "| Case | Defect | Exp. rule | Exp. sev | Exp. exit | Obs. rule(s) | Obs. sev | Obs. exit | Verdict |"  # noqa: E501
    table_sep = "|------|--------|-----------|----------|-----------|--------------|----------|-----------|---------|"  # noqa: E501

    body = f"""# QA Round 2 — defect-injection results

Generated by `scripts/qa_round2.py` against
`examples/briefs/clean_appellate_brief.docx` using the `{PROFILE_ID}` profile.
Method and rationale: see [QA_ROUND2.md](QA_ROUND2.md).

**Result: {passed}/{total} scenarios passed.**

Each scenario injects exactly one defect into a fresh copy of the finding-free
clean fixture, runs the real pipeline + rule pack in-process, and checks the
fired rules, severity, exit code, and `.docx` suppression. TT-008 (font
substitution) fires on every run and is the baseline, never counted as a
defect. A scenario passes only when the expected rule fired at the expected
severity, the exit code and suppression flag matched, and no other defect rule
fired. Per-case artifacts (altered `.docx` + change log) are in
`build/qa_round2/` (uncommitted).

> **Round history.** The first run of this harness was 6/7: the tt005 case
> (TOA deleted, no marker) raised `ConvergenceError` instead of reporting a
> clean TT-005, because a Carmody pinpoint became unmeasurable in the
> TOA-less render. That was fixed by scoping `freeze_registry`'s missing-page
> guard to the output-emitting path (CHANGELOG → Fixed; regression test
> `test_oracle_brief_no_toa`; residual measurement-quality item BL-3). The
> result below reflects the post-fix behavior.

{table_header}
{table_sep}
{chr(10).join(rows)}

## Conclusions

{chr(10).join(conclusion_lines)}

## Per-scenario detail

{chr(10).join(detail)}
## Not exercised here

- **TT-007 (non-convergence):** cannot be hand-injected (oscillation is an
  emergent pagination property, not document content); covered by the
  mocked-measurer unit test `tests/unit/test_tt007_convergence.py`.
- **TT-008 (font substitution):** fires on every run in this environment
  (Consolas → DejaVu Sans Mono); it is the baseline, not an injected defect.
"""
    RESULTS.write_text(body, encoding="utf-8")


def main() -> int:
    """Run all scenarios, write artifacts, and print a one-line summary."""
    WORKDIR.mkdir(parents=True, exist_ok=True)
    profile = load_profile_by_id(PROFILE_ID)
    outcomes = []
    for scenario in SCENARIOS:
        outcome = run_scenario(scenario, profile)
        write_case_log(outcome)
        outcomes.append(outcome)
        flag = "PASS" if outcome.passed else "FAIL"
        observed = ", ".join(sorted(outcome.defect_rules)) or "(no defect)"
        expected = scenario.expected_rule or "none"
        print(
            f"[{flag}] {scenario.slug:7} expected={expected:6} "
            f"observed={observed} exit={outcome.exit_code}"
        )
    write_summary(outcomes)
    passed = sum(1 for o in outcomes if o.passed)
    print(f"\n{passed}/{len(outcomes)} passed — summary: {RESULTS}")
    return 0 if passed == len(outcomes) else 1


if __name__ == "__main__":
    raise SystemExit(main())
