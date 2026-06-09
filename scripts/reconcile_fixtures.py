"""Phase-3 fixture reconciliation — checked-in, idempotent (PRD §22).

Two transformations on the example briefs, each followed by verification against
a real LibreOffice render. The measured page values are taken from THIS
environment's render and treated as truth (LibreOffice + Liberation/DejaVu font
substitution; ±1 drift at the margins is accepted, per spec §3.1).

1. CLEAN reconciliation
   Rewrite ``clean_appellate_brief.docx``'s input Table of Authorities so it
   exactly mirrors the registry toatool measures: the two jurisdictional
   statutes the body cites (``28 U.S.C. § 1331`` / ``§ 1291``) are added, and
   every page list is rewritten to its measured value. After this the input TOA
   diffs clean against the generated TOA — no missing entry (TT-002), no phantom
   entry (TT-004), no stale page (TT-003) — so the clean fixture is finding-free.

2. DIRTY passim nudge
   Append a short trailing argument section (defect D4) to
   ``dirty_motion_brief.docx`` that cites *Carmody* on a sixth physical page, so
   *Carmody* renders as "passim" end-to-end. *Delgado* stays a five-page list,
   preserving the D4-vs-D5 boundary (D5). The dirty fixture's other defects
   (stale pages, the missing *Okafor* entry, the *Ellison* orphan) are left
   intact — it is the defect battery and is meant to produce findings.

Both edits are idempotent: clean's TOA region is detected by placement and
rewritten wholesale every run; dirty's section is inserted only when its sentinel
heading is absent. Re-running this script is a no-op once both are in place.

Usage::

    python scripts/reconcile_fixtures.py          # apply + verify both fixtures
    python scripts/reconcile_fixtures.py --check   # verify only; fail if unreconciled
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any

import docx
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from toatool.engine.profile_loader import load_profile_by_id
from toatool.pipeline import convergence, parser, placement, toa_builder
from toatool.pipeline.input_toa_diff import DiffBaseline, build_baseline

REPO_ROOT = Path(__file__).resolve().parent.parent
BRIEFS = REPO_ROOT / "examples" / "briefs"
CLEAN = BRIEFS / "clean_appellate_brief.docx"
DIRTY = BRIEFS / "dirty_motion_brief.docx"

# Sixth-page Carmody content for the dirty fixture (defect D4). Each paragraph
# cites Carmody once and none cites Delgado, so Carmody gains a sixth physical
# page while Delgado stays at five. Sentinel = the section heading text.
DIRTY_SECTION_HEADING = "V. Defendant's Sufficiency Argument Fails"
DIRTY_SECTION_BODY = (
    "Defendant's final contention regarding the sufficiency of plaintiff's "
    "evidence collapses under settled law. This Court has squarely held that a "
    "plaintiff need not produce direct proof of subjective intent to survive "
    "summary judgment. Carmody, 512 F.3d at 1046.",
    "The objective inquiry that Carmody prescribes does not turn on the "
    "defendant's self-serving characterization of its own conduct. Carmody, "
    "512 F.3d at 1047. A reasonable factfinder could readily conclude that the "
    "communication at issue was misleading under that standard.",
    "Nothing in defendant's submission distinguishes this record from the one "
    "the court confronted in Carmody, 512 F.3d at 1048, where comparable "
    "boilerplate was held insufficient as a matter of law. The same result "
    "follows here.",
    "Defendant's remaining authorities predate Carmody and cannot be reconciled "
    "with its holding. Carmody, 512 F.3d at 1050. To the extent any tension "
    "exists, the controlling circuit precedent governs and forecloses "
    "defendant's position.",
    "Even assuming defendant could identify some residual factual dispute, that "
    "dispute would be immaterial under the governing legal framework. Carmody, "
    "512 F.3d at 1051. Summary judgment for the defendant is therefore "
    "unwarranted on this record.",
    "In short, every strand of defendant's sufficiency argument has already "
    "been considered and rejected by this Court. Carmody, 512 F.3d at 1052. The "
    "motion fails for this independent reason as well.",
)

# Whitespace between an entry's text and its page list. Three spaces is
# unambiguous for the input-TOA diff regex even when a citation ends in digits
# (e.g. "28 U.S.C. § 1331   1").
_ENTRY_GAP = "   "


class ReconcileError(Exception):
    """Raised when a fixture fails its post-edit verification."""


def _style_object(document: DocxDocument, name: str) -> Any:
    """Return a style *object* used by some paragraph, by name, or ``None``.

    These fixtures reference built-in styles that python-docx cannot resolve by
    name (``KeyError``); reusing an existing paragraph's style object sidesteps
    the name lookup (the same technique as ``inserter._make_heading``). Returned
    as ``Any`` because python-docx's style seam is untyped for our purposes.
    """
    for para in document.paragraphs:
        if para.style is not None and para.style.name == name:
            return para.style
    return None


def _delete_paragraph(paragraph: Paragraph) -> None:
    """Remove a paragraph from its document."""
    element = paragraph._p
    element.getparent().remove(element)


def _insert_before(anchor: Paragraph, text: str, style: Any) -> Paragraph:
    """Insert a paragraph carrying ``text``/``style`` immediately before ``anchor``."""
    para = anchor.insert_paragraph_before(text)
    if style is not None:
        para.style = style
    return para


# --------------------------------------------------------------------------- #
# Clean reconciliation
# --------------------------------------------------------------------------- #


def _measure_clean() -> convergence.GenerationResult:
    """Run the generation pipeline on the clean brief."""
    return convergence.generate(CLEAN, load_profile_by_id("frap"))


def _rewrite_clean_toa(result: convergence.GenerationResult) -> None:
    """Replace clean's input-TOA region with one mirroring the measured registry.

    The region (the paragraphs below the ``TABLE OF AUTHORITIES`` heading, up to
    the next heading) is deleted and rebuilt: one bold-free label paragraph per
    non-empty group, then one entry paragraph per authority in profile group and
    sort order, each reading ``<display_full><gap><measured page list>``.
    """
    profile = load_profile_by_id("frap")
    parsed = parser.parse(CLEAN)
    place = placement.detect_placement(parsed, profile.heading.detection_variants)
    if place.mechanism != "heading" or place.heading_match is None:
        raise ReconcileError("clean brief: expected heading placement for its TOA")
    match = place.heading_match

    document = docx.Document(str(CLEAN))
    label_style = _style_object(document, "First Paragraph")
    entry_style = _style_object(document, "Source Code")

    # The paragraph just past the region (the next heading) anchors insertion; it
    # survives the region deletion, so new entries land between heading and it.
    anchor = document.paragraphs[match.region_end]
    for para in document.paragraphs[match.region_start : match.region_end]:
        _delete_paragraph(para)

    toa = toa_builder.build_toa(result.authorities, profile)
    for group in toa.groups:
        _insert_before(anchor, group.label, label_style)
        for authority in group.authorities:
            line = (
                f"{authority.display_full}{_ENTRY_GAP}"
                f"{toa_builder.page_text(authority, profile)}"
            )
            _insert_before(anchor, line, entry_style)

    document.save(str(CLEAN))


def _registry_pages(result: convergence.GenerationResult) -> dict[str, list[int]]:
    """Map authority id → measured page list from a frozen registry."""
    return {auth.authority_id: list(auth.pages) for auth in result.registry.authorities}


def _registry_passim(result: convergence.GenerationResult) -> dict[str, bool]:
    """Map authority id → passim flag from a frozen registry."""
    return {auth.authority_id: auth.passim for auth in result.registry.authorities}


def _verify_clean_finding_free(result: convergence.GenerationResult) -> None:
    """Assert the input TOA diffs clean against the registry (no TT-002/003/004).

    This is the diff-level equivalent of "finding-free": the same identity- and
    page-comparison the rules engine performs, checked before that engine exists.
    """
    profile = load_profile_by_id("frap")
    parsed = parser.parse(CLEAN)
    place = placement.detect_placement(parsed, profile.heading.detection_variants)
    baseline: DiffBaseline = build_baseline(parsed, place, profile)
    if not baseline.available:
        raise ReconcileError("clean brief: no input-TOA baseline after rewrite")

    reg_pages = _registry_pages(result)
    reg_passim = _registry_passim(result)
    seen: set[str] = set()
    for entry in baseline.entries:
        if not entry.parsed or entry.authority_id is None:
            raise ReconcileError(
                f"clean TOA entry did not resolve to an authority: "
                f"{entry.verbatim_text!r}"
            )
        if entry.authority_id not in reg_pages:
            raise ReconcileError(
                f"phantom clean TOA entry (TT-004): {entry.authority_id} "
                f"({entry.verbatim_text!r}) is not in the measured registry"
            )
        want_passim = reg_passim[entry.authority_id]
        if want_passim:
            if not entry.is_passim:
                raise ReconcileError(
                    f"stale clean TOA entry (TT-003): {entry.authority_id} should "
                    f"render passim"
                )
        elif entry.input_pages != reg_pages[entry.authority_id]:
            raise ReconcileError(
                f"stale clean TOA pages (TT-003): {entry.authority_id} has "
                f"{entry.input_pages}, measured {reg_pages[entry.authority_id]}"
            )
        seen.add(entry.authority_id)

    missing = set(reg_pages) - seen
    if missing:
        raise ReconcileError(
            f"missing clean TOA entries (TT-002): {sorted(missing)} are cited in "
            f"the body but absent from the input TOA"
        )


def reconcile_clean(*, apply: bool) -> convergence.GenerationResult:
    """Reconcile (or, with ``apply=False``, only verify) the clean fixture."""
    before = _measure_clean()
    if apply:
        _rewrite_clean_toa(before)
    after = _measure_clean()
    if after.iteration_count != 1 or not after.converged:
        raise ReconcileError(
            f"clean brief must converge on the first check; got "
            f"iterations={after.iteration_count} converged={after.converged}"
        )
    if _registry_pages(before) != _registry_pages(after):
        raise ReconcileError(
            "clean brief measured pages changed after the TOA rewrite — the input "
            "TOA must not affect the generated output's pagination"
        )
    _verify_clean_finding_free(after)
    return after


# --------------------------------------------------------------------------- #
# Dirty passim nudge
# --------------------------------------------------------------------------- #


def _dirty_has_section(document: DocxDocument) -> bool:
    """Whether the dirty fixture already carries the D4 sixth-page section."""
    return any(p.text.strip() == DIRTY_SECTION_HEADING for p in document.paragraphs)


def _insert_dirty_section() -> None:
    """Insert the trailing Carmody section before the CONCLUSION heading."""
    document = docx.Document(str(DIRTY))
    if _dirty_has_section(document):
        return
    heading_style = _style_object(document, "Heading 2")
    body_style = _style_object(document, "Body Text")
    conclusion = next(
        (p for p in document.paragraphs if p.text.strip() == "CONCLUSION"), None
    )
    if conclusion is None:
        raise ReconcileError("dirty brief: CONCLUSION heading not found")

    _insert_before(conclusion, DIRTY_SECTION_HEADING, heading_style)
    for sentence in DIRTY_SECTION_BODY:
        _insert_before(conclusion, sentence, body_style)
    document.save(str(DIRTY))


def _verify_dirty(result: convergence.GenerationResult) -> None:
    """Assert the dirty fixture's D4/D5 boundary and unchanged defects hold."""
    profile = load_profile_by_id("frap")
    by_id = {a.authority_id: a for a in result.authorities}

    carmody = by_id.get("case:f3d:512:1042")
    if carmody is None or not toa_builder.is_passim(carmody, profile):
        pages = carmody.pages() if carmody else None
        raise ReconcileError(f"dirty Carmody must render passim (D4); pages={pages}")
    if len(carmody.pages()) <= profile.passim.threshold_pages:
        raise ReconcileError(
            f"dirty Carmody must span more than {profile.passim.threshold_pages} "
            f"pages; got {carmody.pages()}"
        )

    delgado = by_id.get("case:us:487:213")
    if delgado is None or toa_builder.is_passim(delgado, profile):
        pages = delgado.pages() if delgado else None
        raise ReconcileError(
            f"dirty Delgado must remain a page list, not passim (D5); pages={pages}"
        )

    if "case:f3d:891:655" not in by_id:
        raise ReconcileError("dirty Okafor (D1) disappeared after the edit")
    if "statute:usc:15:1692e" not in by_id:
        raise ReconcileError("dirty § 1692e recognizer entry disappeared")
    if not any("740 F.3d" in occ.raw_text for occ in result.unresolved):
        raise ReconcileError("dirty Ellison orphan (D3) disappeared after the edit")
    if not result.converged:
        raise ReconcileError("dirty brief failed to converge after the edit")


def reconcile_dirty(*, apply: bool) -> convergence.GenerationResult:
    """Reconcile (or, with ``apply=False``, only verify) the dirty fixture."""
    if apply:
        _insert_dirty_section()
    result = convergence.generate(DIRTY, load_profile_by_id("frap"))
    _verify_dirty(result)
    return result


# --------------------------------------------------------------------------- #
# Reporting
# --------------------------------------------------------------------------- #


def _print_truth_table(name: str, result: convergence.GenerationResult) -> None:
    """Print the measured authority → pages table (for FIXTURES_README upkeep)."""
    profile = load_profile_by_id("frap")
    by_id = {a.authority_id: a for a in result.authorities}
    print(f"\n  {name} — measured authorities (this environment):")
    for auth in result.registry.authorities:
        working = by_id[auth.authority_id]
        rendered = toa_builder.page_text(working, profile)
        print(f"    {auth.display_full:<70} {rendered}")
    meta = result.registry.run_metadata
    subs = ", ".join(f"{s.original} → {s.substitute}" for s in meta.font_substitutions)
    print(
        f"    [render {meta.render_engine} {meta.render_engine_version}; "
        f"iterations={meta.iteration_count}; converged={meta.converged}; "
        f"font subs: {subs or 'none'}]"
    )


def main(argv: list[str] | None = None) -> int:
    """Apply (or verify) both fixture reconciliations and print measured truth."""
    arg_parser = argparse.ArgumentParser(description=__doc__)
    arg_parser.add_argument(
        "--check",
        action="store_true",
        help="verify only; do not modify the fixtures",
    )
    args = arg_parser.parse_args(argv)
    apply = not args.check

    mode = "Verifying" if args.check else "Reconciling"
    print(f"{mode} example fixtures (LibreOffice render; this may take ~15s)...")
    try:
        clean = reconcile_clean(apply=apply)
        print("  clean: input TOA mirrors the measured registry; finding-free ✓")
        dirty = reconcile_dirty(apply=apply)
        print("  dirty: Carmody renders passim end-to-end; Delgado stays a list ✓")
    except ReconcileError as exc:
        print(f"\nFAILED: {exc}", file=sys.stderr)
        return 1

    _print_truth_table("clean_appellate_brief.docx", clean)
    _print_truth_table("dirty_motion_brief.docx", dirty)
    print("\nDone. Update examples/briefs/README.md tables from the values above.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
