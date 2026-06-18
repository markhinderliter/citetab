"""The missing-LibreOffice path must fail gracefully, instructively, and in order.

LibreOffice is a required render dependency (page measurement). When it is absent
or misconfigured, ``find_libreoffice`` raises ``RenderError``. Both surfaces must
turn that into a clean FAILED outcome (exit 2) carrying the install instructions —
never an unhandled exception or a stack trace.

These tests simulate the absence by monkeypatching ``renderer.find_libreoffice``
to raise the real not-found ``RenderError`` (one target covers both the render
path's internal call and core's re-probe); they never require LibreOffice to be
actually missing on the host, so they run anywhere with no real render. They lock
in two things a future refactor must not regress — both invisible in CI, where
LibreOffice is installed:

1. **Graceful + framed (option 3):** a valid file on a LibreOffice-less machine
   fails cleanly (exit 2) with the install message, NOT a "couldn't process that
   file" prefix (a missing dependency is an environment problem, not a bad file).
2. **Ordering (option B):** when the input is *also* bad/empty, the FILE error
   wins — the parse/empty-document checks run before the render probe, so the user
   learns their file is unusable (e.g. "run OCR") instead of being sent to install
   a 300MB dependency they would then discover did not fix the file.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest
from click.testing import CliRunner

from citetab.cli import main as cli_main
from citetab.core import Outcome, run_generation
from citetab.pipeline import renderer
from citetab.pipeline.renderer import RenderError

#: The genuine not-found message find_libreoffice raises (renderer.py).
_NOT_FOUND = (
    "LibreOffice was not found. It is a required system dependency for measuring "
    "page numbers. Install it from https://www.libreoffice.org/download/ (see the "
    "README 'System requirements' section), or set the CITETAB_LIBREOFFICE "
    "environment variable to the full path of the 'soffice' executable, and try "
    "again."
)


def _raise_not_found() -> str:
    raise RenderError(_NOT_FOUND)


def _valid_brief(tmp_path: Path) -> Path:
    """A genuinely valid, non-empty .docx — parses and has readable body text, so
    it gets past the parse/empty guards and reaches the render probe."""
    brief = tmp_path / "brief.docx"
    document = docx.Document()
    document.add_paragraph("The court considered Roe v. Wade, 410 U.S. 113 (1973).")
    document.save(str(brief))
    return brief


# --------------------------------------------------------------------------- #
# Graceful + environment-framed (option 3): valid file, LibreOffice absent
# --------------------------------------------------------------------------- #


def test_run_generation_missing_libreoffice_is_graceful(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """A valid brief + absent LibreOffice → clean FAILED, environment-framed."""
    monkeypatch.setattr(renderer, "find_libreoffice", _raise_not_found)
    brief = _valid_brief(tmp_path)

    result = run_generation(brief)

    assert result.outcome is Outcome.FAILED
    assert result.exit_code == 2
    assert result.docx_path is None
    assert result.report_path is None
    # Instructive: points at installing the dependency.
    assert "libreoffice.org" in result.message
    assert "libreoffice.org" in (result.error or "")
    # Framing (option 3): a missing dependency is not a bad file.
    assert "couldn't process that file" not in result.message.lower()


def test_cli_generate_missing_libreoffice_exits_2_cleanly(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """The CLI surfaces the missing dependency with exit 2, no traceback."""
    monkeypatch.setattr(renderer, "find_libreoffice", _raise_not_found)
    brief = _valid_brief(tmp_path)

    runner = CliRunner()
    invoked = runner.invoke(cli_main, ["generate", str(brief)])

    assert invoked.exit_code == 2, invoked.output
    assert invoked.exception is None or isinstance(invoked.exception, SystemExit)
    assert "libreoffice.org" in invoked.output
    assert "couldn't process that file" not in invoked.output.lower()


# --------------------------------------------------------------------------- #
# Ordering (option B): bad/empty input wins over missing LibreOffice
# --------------------------------------------------------------------------- #


def test_empty_document_error_wins_over_missing_libreoffice(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """A textless brief + absent LibreOffice → the FILE (OCR) error wins."""
    monkeypatch.setattr(renderer, "find_libreoffice", _raise_not_found)
    textless = tmp_path / "scanned.docx"
    docx.Document().save(str(textless))  # valid .docx, no readable text

    result = run_generation(textless)

    assert result.outcome is Outcome.FAILED
    assert result.exit_code == 2
    lower = result.message.lower()
    # The file problem is reported, framed as such…
    assert "no readable text" in lower
    assert "ocr" in lower
    assert "couldn't process that file" in lower
    # …NOT the missing-dependency message.
    assert "libreoffice.org" not in result.message
    assert "libreoffice was not found" not in lower


def test_parser_error_wins_over_missing_libreoffice(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """A non-.docx input + absent LibreOffice → the FILE (parse) error wins."""
    monkeypatch.setattr(renderer, "find_libreoffice", _raise_not_found)
    bad = tmp_path / "not_really.docx"
    bad.write_text("this is plain text, not a docx archive", encoding="utf-8")

    result = run_generation(bad)

    assert result.outcome is Outcome.FAILED
    assert result.exit_code == 2
    # File problem framing, not the dependency message.
    assert "couldn't process that file" in result.message.lower()
    assert "libreoffice.org" not in result.message
