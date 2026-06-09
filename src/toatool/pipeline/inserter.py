"""Insert the generated TOA into a document copy (FR-02, FR-09).

The input is never modified; this operates on a working copy. Placement follows
the detected mechanism:

- **heading / flag:** keep the matched heading paragraph, delete the old region
  beneath it, and insert the generated entries there.
- **marker:** replace the ``[[TOA]]`` marker with the standard heading plus the
  generated entries, and consume the marker. The generated heading carries the
  profile's standard text so a re-run detects placement via the heading path
  (the idempotency handoff, spec §2.4).

The TOA is written as **static formatted content** — never Word TA/TOA field
codes. Each entry keeps a reference to its page-number run so the convergence
loop can rewrite just the digits each iteration without changing the structure.
"""

from __future__ import annotations

from dataclasses import dataclass, field

from docx.document import Document as DocxDocument
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from toatool.engine.profile_loader import CourtProfile
from toatool.pipeline.placement import PlacementResult
from toatool.pipeline.toa_builder import ToaModel, page_text
from toatool.pipeline.working import WorkingAuthority

_PAGE_TAB_INCHES = 6.5


@dataclass
class InsertedToa:
    """Handle to the inserted TOA: each entry's page run, for live updates."""

    entries: list[tuple[WorkingAuthority, Run]] = field(default_factory=list)


def _move_after(element: object, anchor: object) -> None:
    """Move ``element`` to sit immediately after ``anchor`` in the XML tree."""
    parent = element.getparent()  # type: ignore[attr-defined]
    parent.remove(element)
    anchor.addnext(element)  # type: ignore[attr-defined]


def _new_paragraph(document: DocxDocument) -> Paragraph:
    """Create a paragraph (appended at the document end) for later relocation."""
    return document.add_paragraph()


def _make_heading(document: DocxDocument, text: str) -> Paragraph:
    """Create a level-1 heading paragraph, reusing the document's heading style.

    Assigning an existing style *object* avoids python-docx's name-based lookup,
    which raises for built-in styles that a document references but does not
    define under that exact name. The standard heading style is required so a
    re-run detects the generated TOA via the heading path (idempotency handoff).
    """
    style = None
    for para in document.paragraphs:
        if para.style is not None and para.style.name == "Heading 1":
            style = para.style
            break
    if style is None:
        try:
            style = document.styles["Heading 1"]
        except KeyError:
            style = None
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    if style is not None:
        paragraph.style = style
    else:
        run.bold = True
    return paragraph


def _case_name_split(authority: WorkingAuthority) -> tuple[str, str]:
    """Split a case display into the italic name part and the roman remainder."""
    plaintiff = authority.components.get("plaintiff")
    defendant = authority.components.get("defendant")
    if plaintiff and defendant:
        name = f"{plaintiff} v. {defendant}"
    elif plaintiff:
        name = str(plaintiff)
    else:
        return "", authority.display_full
    if authority.display_full.startswith(name):
        return name, authority.display_full[len(name) :]
    return "", authority.display_full


def _fill_entry(
    paragraph: Paragraph,
    authority: WorkingAuthority,
    profile: CourtProfile,
) -> Run:
    """Populate an entry paragraph and return its (updatable) page-number run."""
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(_PAGE_TAB_INCHES),
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.DOTS if profile.formatting.dot_leader else WD_TAB_LEADER.SPACES,
    )
    if authority.type == "case" and profile.formatting.italicize_case_names:
        italic_part, roman_part = _case_name_split(authority)
        if italic_part:
            run = paragraph.add_run(italic_part)
            run.italic = True
        paragraph.add_run(roman_part)
    else:
        paragraph.add_run(authority.display_full)
    paragraph.add_run("\t")
    return paragraph.add_run(page_text(authority, profile))


def _delete_paragraph(paragraph: Paragraph) -> None:
    """Remove a paragraph from its document."""
    element = paragraph._p
    element.getparent().remove(element)


def insert_toa(
    document: DocxDocument,
    placement: PlacementResult,
    toa: ToaModel,
    profile: CourtProfile,
) -> InsertedToa:
    """Insert the generated TOA into ``document`` per the placement result.

    Args:
        document: The working document copy to modify in place.
        placement: The detected placement (heading/flag/marker).
        toa: The grouped, sorted TOA structure.
        profile: The active court profile.

    Returns:
        A handle whose ``entries`` map each authority to its page-number run for
        per-iteration updates.

    Raises:
        ValueError: If called with a ``none`` placement (the TT-005 path writes
            no document; the caller must not reach here).
    """
    if placement.mechanism == "none":
        raise ValueError("cannot insert a TOA when no placement was found")

    paragraphs = document.paragraphs
    if placement.mechanism == "marker":
        assert placement.marker_index is not None
        anchor_para = paragraphs[placement.marker_index]
        heading = _make_heading(document, profile.heading.generated_text)
        _move_after(heading._p, anchor_para._p)
        cursor = heading._p
        marker_to_delete: Paragraph | None = anchor_para
    else:
        assert placement.heading_match is not None
        match = placement.heading_match
        for para in paragraphs[match.region_start : match.region_end]:
            _delete_paragraph(para)
        cursor = paragraphs[match.heading_index]._p
        marker_to_delete = None

    inserted = InsertedToa()
    for group in toa.groups:
        label_para = _new_paragraph(document)
        label_run = label_para.add_run(group.label)
        label_run.bold = True
        _move_after(label_para._p, cursor)
        cursor = label_para._p
        for authority in group.authorities:
            entry_para = _new_paragraph(document)
            page_run = _fill_entry(entry_para, authority, profile)
            _move_after(entry_para._p, cursor)
            cursor = entry_para._p
            inserted.entries.append((authority, page_run))

    if marker_to_delete is not None:
        _delete_paragraph(marker_to_delete)

    return inserted


def update_pages(inserted: InsertedToa, profile: CourtProfile) -> None:
    """Rewrite each entry's page-number run from its authority's current pages."""
    for authority, run in inserted.entries:
        run.text = page_text(authority, profile)
