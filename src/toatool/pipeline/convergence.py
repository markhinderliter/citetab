"""The fixed-point regeneration orchestrator (spec §3.2, FR-08).

Parse once, build once, then measure → write → re-render → check, repeating only
the check until no occurrence's page changes, capped at 5 iterations. The
canonical :class:`~toatool.models.registry.CitationRegistry` is frozen exactly
once, *after* the loop settles: pages, iteration count, and font substitutions
are accumulated in mutable working structures during the loop and converted in a
single step. The registry is never mutated mid-loop.

A hard invariant guards the freeze: every occurrence that feeds a TOA entry must
have a measured page. A converged registry with a null page is a bug, not a
silently emitted entry — so it raises rather than producing an entry with a
missing page.
"""

from __future__ import annotations

import hashlib
import tempfile
from dataclasses import dataclass
from pathlib import Path

import docx
from docx.document import Document as DocxDocument

from toatool import __version__
from toatool.engine.profile_loader import CourtProfile
from toatool.engine.rule_loader import RULE_PACK_VERSION
from toatool.models.registry import (
    Authority,
    CitationRegistry,
    FontSubstitution,
    Occurrence,
    RunMetadata,
)
from toatool.pipeline import (
    extractor,
    inserter,
    locator,
    parser,
    placement,
    renderer,
    resolver,
    toa_builder,
)
from toatool.pipeline.working import WorkingAuthority, WorkingOccurrence

ITERATION_CAP = 5


class ConvergenceError(Exception):
    """Raised when the converged result violates a hard invariant."""


@dataclass
class GenerationResult:
    """The outcome of running the generation pipeline on one brief."""

    registry: CitationRegistry
    placement: placement.PlacementResult
    working_document: DocxDocument
    inserted: inserter.InsertedToa | None
    authorities: list[WorkingAuthority]
    unresolved: list[WorkingOccurrence]
    toa: toa_builder.ToaModel
    iteration_count: int
    converged: bool
    page_history: list[dict[str, list[int]]]
    blocks_docx_output: bool


def _excluded_and_anchor(
    parsed: parser.ParsedDocument, place: placement.PlacementResult
) -> tuple[set[int], str | None]:
    """Return the paragraph indices to exclude from the body, and the TOA anchor."""
    if place.mechanism in ("heading", "flag") and place.heading_match is not None:
        match = place.heading_match
        excluded = set(range(match.heading_index, match.region_end))
        anchor = (
            parsed.paragraphs[match.region_end].text
            if match.region_end < len(parsed.paragraphs)
            else None
        )
        return excluded, anchor
    if place.mechanism == "marker" and place.marker_index is not None:
        excluded = {place.marker_index}
        following = place.marker_index + 1
        anchor = (
            parsed.paragraphs[following].text
            if following < len(parsed.paragraphs)
            else None
        )
        return excluded, anchor
    return set(), None


def _ordered_occurrences(
    authorities: list[WorkingAuthority], unresolved: list[WorkingOccurrence]
) -> list[WorkingOccurrence]:
    """All occurrences (authority + orphan) in document order."""
    occurrences = [occ for auth in authorities for occ in auth.occurrences]
    occurrences.extend(unresolved)
    occurrences.sort(key=lambda o: (o.paragraph_index, o.char_span))
    return occurrences


def _page_snapshot(authorities: list[WorkingAuthority]) -> dict[str, list[int]]:
    """Capture each authority's current pages, for convergence comparison/history."""
    return {auth.authority_id: auth.pages() for auth in authorities}


def _input_sha256(path: Path) -> str:
    """Return the SHA-256 hex digest of the input file."""
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _run_convergence_loop(
    working: DocxDocument,
    inserted: inserter.InsertedToa,
    profile: CourtProfile,
    occurrences: list[WorkingOccurrence],
    authorities: list[WorkingAuthority],
    anchor: str | None,
    workdir: Path,
) -> tuple[int, bool, list[dict[str, list[int]]]]:
    """Run the measure/write/check loop; return (iterations, converged, history)."""
    working_path = workdir / "working.docx"
    written: dict[str, list[int]] | None = None
    history: list[dict[str, list[int]]] = []
    iterations = 0
    converged = False

    for attempt in range(1, ITERATION_CAP + 1):
        working.save(str(working_path))
        pdf_path = renderer.render_to_pdf(working_path, workdir)
        locator.locate_occurrences(pdf_path, occurrences, anchor)
        snapshot = _page_snapshot(authorities)
        history.append(snapshot)
        if snapshot == written:
            converged = True
            break
        written = snapshot
        inserter.update_pages(inserted, profile)
        iterations = attempt

    return iterations, converged, history


def freeze_registry(
    authorities: list[WorkingAuthority],
    profile: CourtProfile,
    *,
    input_path: Path,
    render_engine: str,
    render_version: str,
    font_substitutions: list[renderer.FontSubstitution],
    iteration_count: int,
    converged: bool,
) -> CitationRegistry:
    """Freeze the canonical registry from settled working authorities.

    An occurrence the locator could not place keeps ``page = None`` rather than
    being fatal: the registry is frozen as-is, the ``.docx`` is still written with
    the pages that *were* measured, and TT-009 discloses the unmeasured
    occurrence (rendered ``p.?``). This is honest degradation over a crash or a
    silently dropped page — the BL-3 fix found in QA Round 3 (see
    ``docs/QA_ROUND3.md``). Earlier builds raised ``ConvergenceError`` here.

    Args:
        authorities: The working authorities, with pages assigned where measured.
        profile: The active court profile.
        input_path: The input brief (for its hash).
        render_engine: Render engine identity.
        render_version: Render engine version.
        font_substitutions: Font substitutions detected for the render.
        iteration_count: Iterations the loop took.
        converged: Whether the loop converged within the cap.

    Returns:
        The immutable :class:`CitationRegistry`.
    """
    frozen: list[Authority] = []
    for auth in sorted(authorities, key=lambda a: (a.group, a.sort_key)):
        occurrences: list[Occurrence] = []
        for occ in auth.occurrences:
            occurrences.append(
                Occurrence(
                    form=occ.form,
                    raw_text=occ.raw_text,
                    paragraph_index=occ.paragraph_index,
                    char_span=occ.char_span,
                    confidence="high" if occ.confidence == "high" else "medium",
                    page=occ.page,
                    pincite=occ.pincite,
                )
            )
        frozen.append(
            Authority(
                authority_id=auth.authority_id,
                type=auth.type,
                components=auth.components,
                display_full=auth.display_full,
                sort_key=auth.sort_key,
                group=auth.group,
                occurrences=occurrences,
                pages=auth.pages(),
                passim=toa_builder.is_passim(auth, profile),
                display_short=auth.display_short,
            )
        )

    metadata = RunMetadata(
        engine_version=__version__,
        rule_pack_version=RULE_PACK_VERSION,
        profile_id=profile.id,
        profile_version=profile.version,
        render_engine=render_engine,
        render_engine_version=render_version,
        font_substitution_occurred=bool(font_substitutions),
        iteration_count=max(iteration_count, 1),
        converged=converged,
        input_sha256=_input_sha256(input_path),
        font_substitutions=[
            FontSubstitution(original=sub.original, substitute=sub.substitute)
            for sub in font_substitutions
        ],
    )
    return CitationRegistry(authorities=frozen, run_metadata=metadata)


def generate(
    input_path: Path,
    profile: CourtProfile,
    *,
    toa_heading: str | None = None,
    workdir: Path | None = None,
) -> GenerationResult:
    """Run the full generation pipeline on one brief.

    Args:
        input_path: The input ``.docx`` brief (never modified).
        profile: The active court profile.
        toa_heading: An explicit ``--toa-heading`` override, if any.
        workdir: Directory for render scratch files; a temp dir is used if None.

    Returns:
        A :class:`GenerationResult` with the frozen registry, the working
        document (for output writing), and loop diagnostics.
    """
    if workdir is None:
        with tempfile.TemporaryDirectory() as temp:
            return generate(
                input_path, profile, toa_heading=toa_heading, workdir=Path(temp)
            )

    parsed = parser.parse(input_path)
    place = placement.detect_placement(
        parsed, profile.heading.detection_variants, toa_heading
    )
    excluded, anchor = _excluded_and_anchor(parsed, place)
    body = extractor.build_body(parsed, excluded)
    if not body.text.strip():
        raise extractor.EmptyDocumentError(
            f"no readable text found in '{input_path.name}'. If the document is "
            f"scanned or image-only, run OCR to make its text selectable, then "
            f"try again."
        )
    resolved = resolver.resolve(parsed, body, profile)
    toa = toa_builder.build_toa(resolved.authorities, profile)
    occurrences = _ordered_occurrences(resolved.authorities, resolved.unresolved)

    engine_name, engine_version = renderer.render_engine_info()
    font_subs = renderer.detect_font_substitutions(input_path)

    working = docx.Document(str(input_path))
    if place.mechanism == "none":
        iterations, converged, inserted, history = _measure_without_placement(
            working, occurrences, anchor, workdir
        )
        blocks = True
    else:
        inserted = inserter.insert_toa(working, place, toa, profile)
        iterations, converged, history = _run_convergence_loop(
            working,
            inserted,
            profile,
            occurrences,
            resolved.authorities,
            anchor,
            workdir,
        )
        blocks = False

    registry = freeze_registry(
        resolved.authorities,
        profile,
        input_path=input_path,
        render_engine=engine_name,
        render_version=engine_version,
        font_substitutions=font_subs,
        iteration_count=iterations,
        converged=converged,
    )

    return GenerationResult(
        registry=registry,
        placement=place,
        working_document=working,
        inserted=inserted,
        authorities=resolved.authorities,
        unresolved=resolved.unresolved,
        toa=toa,
        iteration_count=iterations,
        converged=converged,
        page_history=history,
        blocks_docx_output=blocks,
    )


def _measure_without_placement(
    working: DocxDocument,
    occurrences: list[WorkingOccurrence],
    anchor: str | None,
    workdir: Path,
) -> tuple[int, bool, inserter.InsertedToa | None, list[dict[str, list[int]]]]:
    """Measure occurrence pages on the unmodified document (the TT-005 path).

    No TOA is inserted (there is nowhere to put it), so a single render measures
    body pages for the report's generated TOA. Returns one iteration, converged.
    """
    working_path = workdir / "noplacement.docx"
    working.save(str(working_path))
    pdf_path = renderer.render_to_pdf(working_path, workdir)
    locator.locate_occurrences(pdf_path, occurrences, anchor)
    return 1, True, None, []
